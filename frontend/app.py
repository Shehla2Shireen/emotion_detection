import streamlit as st
import cv2
import io
import numpy as np
import requests
from collections import deque, defaultdict
import json
import os
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor

# =========================
# CONFIG
# =========================
BACKEND_URL = "http://localhost:8000/predict"
ROLLING_WINDOW = 15
FPS = 15
SETTINGS_FILE = "admin_settings.json"

EMOTIONS = ['Angry','Happy','Sad','Surprise','Neutral','Disgust','Fear']

# Ideal defaults (for reset)
IDEAL_EMOTION_WEIGHTS = {
    'Angry': 0.29,
    'Sad': 0.25,
    'Fear': 0.33,
    'Disgust': 0.13,
    # Non-negative emotions get weight 0 (ignored in stress calc)
    'Happy': 0.0,
    'Surprise': 0.0,
    'Neutral': 0.0
}
IDEAL_STRESS_WEIGHTS = {'emotions': 0.7, 'eye_contact': 0.3}  # sum=1
IDEAL_EXPECTED_RANGES = {
    "Neutral": (0.50, 0.70),
    "Happy": (0.20, 0.30),
    "Surprise": (0.05, 0.10),
    "Negatives": (0.00, 0.10)
}

# =========================
# SETTINGS PERSISTENCE
# =========================
def load_settings():
    if os.path.exists(SETTINGS_FILE):
        with open(SETTINGS_FILE, "r") as f:
            return json.load(f)
    return {
        "EMOTION_WEIGHTS": IDEAL_EMOTION_WEIGHTS,
        "STRESS_WEIGHTS": IDEAL_STRESS_WEIGHTS,
        "EXPECTED_RANGES": IDEAL_EXPECTED_RANGES
    }

def save_settings():
    settings = {
        "EMOTION_WEIGHTS": st.session_state.EMOTION_WEIGHTS,
        "STRESS_WEIGHTS": st.session_state.STRESS_WEIGHTS,
        "EXPECTED_RANGES": st.session_state.EXPECTED_RANGES
    }
    with open(SETTINGS_FILE, "w") as f:
        json.dump(settings, f)

# =========================
# SESSION STATE DEFAULTS
# =========================
if "history" not in st.session_state:
    st.session_state.history = deque(maxlen=ROLLING_WINDOW)
if "eye_contact_history" not in st.session_state:
    st.session_state.eye_contact_history = deque(maxlen=ROLLING_WINDOW)
if "eye_contact_binary_history" not in st.session_state:
    st.session_state.eye_contact_binary_history = deque(maxlen=ROLLING_WINDOW)
if "stress_history" not in st.session_state:
    st.session_state.stress_history = deque(maxlen=ROLLING_WINDOW)
if "camera_active" not in st.session_state:
    st.session_state.camera_active = False
if "smoothed_stress" not in st.session_state:
    st.session_state.smoothed_stress = 0

# Load settings into session state
loaded_settings = load_settings()
if "EMOTION_WEIGHTS" not in st.session_state:
    st.session_state.EMOTION_WEIGHTS = loaded_settings["EMOTION_WEIGHTS"]
if "STRESS_WEIGHTS" not in st.session_state:
    st.session_state.STRESS_WEIGHTS = loaded_settings["STRESS_WEIGHTS"]
if "EXPECTED_RANGES" not in st.session_state:
    st.session_state.EXPECTED_RANGES = loaded_settings["EXPECTED_RANGES"]

ALPHA = 0.3

# =========================
# HELPERS
# =========================
def get_stress_label(stress_score: float):
    if stress_score < 5:
        return "Low Stress", "🟢"
    elif stress_score < 10:
        return "Moderate Stress", "🟡"
    else:
        return "High Stress", "🔴"

def get_interview_status(stress_score, eye_contact_percentage, evaluation_report):
    """Determine overall interview status based on multiple metrics"""
    stress_score_val = 100 - stress_score  # invert so higher is better
    eye_contact_score = eye_contact_percentage  # percentage of time eye contact was 100%
    compliance_score = 0
    total_ranges = len(evaluation_report)
    for _, (_, ok) in evaluation_report.items():
        compliance_score += 100 if ok else 60
    compliance_score /= total_ranges
    overall_score = (
        stress_score_val * 0.4 + 
        eye_contact_score * 0.3 + 
        compliance_score * 0.3
    )
    if overall_score >= 80:
        return "Good ✅", overall_score, {
            "stress_score": stress_score_val,
            "eye_contact_score": eye_contact_score,
            "emotion_compliance_score": compliance_score
        }
    elif overall_score >= 60:
        return "Average ⚠️", overall_score, {
            "stress_score": stress_score_val,
            "eye_contact_score": eye_contact_score,
            "emotion_compliance_score": compliance_score
        }
    else:
        return "Bad ❌", overall_score, {
            "stress_score": stress_score_val,
            "eye_contact_score": eye_contact_score,
            "emotion_compliance_score": compliance_score
        }

def post_to_backend(img_bgr: np.ndarray):
    _, buf = cv2.imencode(".jpg", img_bgr)
    files = {"file": ("frame.jpg", io.BytesIO(buf.tobytes()), "image/jpeg")}
    r = requests.post(BACKEND_URL, files=files, timeout=30)
    r.raise_for_status()
    return r.json()

def calculate_stress_level(emotions, eye_contact_percentage):
    weighted_sum = 0
    for emo in ['Angry', 'Sad', 'Fear', 'Disgust']:
        weighted_sum += emotions.get(emo, 0) * st.session_state.EMOTION_WEIGHTS.get(emo, 0)
    emotion_component = min(100, weighted_sum * 100)
    eye_contact_component = min(100, (100 - eye_contact_percentage))
    raw_stress = (
        st.session_state.STRESS_WEIGHTS['emotions'] * emotion_component +
        st.session_state.STRESS_WEIGHTS['eye_contact'] * eye_contact_component
    )
    smoothed = ALPHA * raw_stress + (1 - ALPHA) * st.session_state.smoothed_stress
    st.session_state.smoothed_stress = smoothed
    return smoothed

def update_rollups(emotion_result):
    probs = emotion_result.get("all_predictions", {})
    full = {e: 0.0 for e in EMOTIONS}
    for k, v in probs.items():
        full[k] += float(v)
    st.session_state.history.append(full)
    
    # Get eye contact value and convert to binary (0 or 100)
    eye_contact_value = emotion_result.get("eye_contact", 0)
    eye_contact_binary = 100 if eye_contact_value >= 95 else 0  # Consider 95%+ as good eye contact
    
    st.session_state.eye_contact_history.append(eye_contact_value)
    st.session_state.eye_contact_binary_history.append(eye_contact_binary)

    emotion_avg = defaultdict(float)
    for row in st.session_state.history:
        for k, v in row.items():
            emotion_avg[k] += v
    n = max(1, len(st.session_state.history))
    for k in emotion_avg:
        emotion_avg[k] /= n

    # Calculate percentage of time eye contact was 100%
    eye_contact_percentage = sum(st.session_state.eye_contact_binary_history) / max(1, len(st.session_state.eye_contact_binary_history))
    eye_contact_avg = sum(st.session_state.eye_contact_history) / max(1, len(st.session_state.eye_contact_history))

    current_stress = calculate_stress_level(emotion_avg, eye_contact_percentage)
    st.session_state.stress_history.append(current_stress)
    stress_avg = sum(st.session_state.stress_history) / max(1, len(st.session_state.stress_history))
    stress_label, stress_emoji = get_stress_label(stress_avg)

    if eye_contact_percentage >= 80:
        eye_contact_perf = "Good"
    elif eye_contact_percentage >= 40:
        eye_contact_perf = "Average"
    else:
        eye_contact_perf = "Poor"

    # Determine dominant emotion
    dominant_emotion = max(emotion_avg.items(), key=lambda kv: kv[1])[0]

    return emotion_avg, eye_contact_avg, eye_contact_percentage, stress_avg, stress_label, stress_emoji, eye_contact_perf, dominant_emotion

def evaluate_emotion_distribution(emotion_avg):
    total = sum(emotion_avg.values())
    norm = {k: v/total for k,v in emotion_avg.items()} if total > 0 else emotion_avg
    report = {}
    for key, (low, high) in st.session_state.EXPECTED_RANGES.items():
        if key == "Negatives":
            neg_val = norm.get("Angry",0)+norm.get("Sad",0)+norm.get("Fear",0)+norm.get("Disgust",0)
            report[key] = (neg_val, low <= neg_val <= high)
        else:
            val = norm.get(key, 0)
            report[key] = (val, low <= val <= high)
    return report

def add_colored_text(paragraph, text, color):
    run = paragraph.add_run(text)
    if color == "green":
        run.font.color.rgb = RGBColor(0, 128, 0)
    elif color == "red":
        run.font.color.rgb = RGBColor(220, 20, 60)
    elif color == "orange":
        run.font.color.rgb = RGBColor(255, 140, 0)
    return paragraph

def generate_report(emotion_avg, eye_contact_avg, eye_contact_percentage, stress_avg, stress_label, evaluation_report):
    doc = Document()
    doc.add_heading("Interview Analysis Report", 0)
    
    interview_status, overall_score, component_scores = get_interview_status(
        stress_avg, eye_contact_percentage, evaluation_report
    )
    
    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    status_para = doc.add_paragraph()
    status_para.add_run("Overall Interview Status: ")
    if "Good" in interview_status:
        add_colored_text(status_para, interview_status, "green")
    elif "Average" in interview_status:
        add_colored_text(status_para, interview_status, "orange")
    else:
        add_colored_text(status_para, interview_status, "red")
    
    score_para = doc.add_paragraph()
    score_para.add_run("Overall Score: ")
    if overall_score >= 80:
        add_colored_text(score_para, f"{overall_score:.1f}/100", "green")
    elif overall_score >= 60:
        add_colored_text(score_para, f"{overall_score:.1f}/100", "orange")
    else:
        add_colored_text(score_para, f"{overall_score:.1f}/100", "red")
    
    # Ideal parameters reference
    doc.add_heading("Ideal Parameters Reference", level=2)
    ideal_para = doc.add_paragraph()
    ideal_para.add_run("• Stress Level: ").bold = True
    ideal_para.add_run("Below 40/100 (Lower is better)\n")
    ideal_para.add_run("• Eye Contact: ").bold = True
    ideal_para.add_run("70-100% of time with good eye contact (Higher is better)\n")
    ideal_para.add_run("• Emotion Ranges: ").bold = True
    ideal_para.add_run("Neutral (50-70%), Happy (20-30%), Surprise (5-10%), Negative Emotions (0-10%)")
    
    # Detailed Metrics
    doc.add_heading("Detailed Metrics", level=1)
    metrics_table = doc.add_table(rows=4, cols=4)
    metrics_table.style = 'Table Grid'
    
    # Table headers
    metrics_table.cell(0, 0).text = "Metric"
    metrics_table.cell(0, 1).text = "Value"
    metrics_table.cell(0, 2).text = "Ideal Range"
    metrics_table.cell(0, 3).text = "Status"
    
    # Stress
    metrics_table.cell(1, 0).text = "Stress Level"
    metrics_table.cell(1, 1).text = f"{stress_avg:.1f}/100"
    metrics_table.cell(1, 2).text = "0-40"
    stress_cell = metrics_table.cell(1, 3)
    if stress_avg < 40:
        stress_cell.text = "Good ✅"
        for paragraph in stress_cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 128, 0)
    elif stress_avg < 65:
        stress_cell.text = "Average ⚠️"
        for paragraph in stress_cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 140, 0)
    else:
        stress_cell.text = "Poor ❌"
        for paragraph in stress_cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(220, 20, 60)
    
    # Eye Contact (Percentage of time with good eye contact)
    metrics_table.cell(2, 0).text = "Eye Contact Time"
    metrics_table.cell(2, 1).text = f"{eye_contact_percentage:.1f}%"
    metrics_table.cell(2, 2).text = "70-100%"
    eye_contact_cell = metrics_table.cell(2, 3)
    if eye_contact_percentage >= 70:
        eye_contact_cell.text = "Good ✅"
        for paragraph in eye_contact_cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 128, 0)
    elif eye_contact_percentage >= 40:
        eye_contact_cell.text = "Average ⚠️"
        for paragraph in eye_contact_cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 140, 0)
    else:
        eye_contact_cell.text = "Poor ❌"
        for paragraph in eye_contact_cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(220, 20, 60)
    
    # Emotion Compliance
    compliance_score = component_scores["emotion_compliance_score"]
    metrics_table.cell(3, 0).text = "Emotion Compliance"
    metrics_table.cell(3, 1).text = f"{compliance_score:.1f}%"
    metrics_table.cell(3, 2).text = "80-100%"
    compliance_cell = metrics_table.cell(3, 3)
    if compliance_score >= 80:
        compliance_cell.text = "Good ✅"
        for paragraph in compliance_cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 128, 0)
    elif compliance_score >= 60:
        compliance_cell.text = "Average ⚠️"
        for paragraph in compliance_cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 140, 0)
    else:
        compliance_cell.text = "Poor ❌"
        for paragraph in compliance_cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(220, 20, 60)

    # Component Scores Breakdown
    doc.add_heading("Component Scores Breakdown", level=1)
    component_table = doc.add_table(rows=4, cols=3)
    component_table.style = 'Table Grid'
    
    # Table headers
    component_table.cell(0, 0).text = "Component"
    component_table.cell(0, 1).text = "Score"
    component_table.cell(0, 2).text = "Weight"
    
    # Stress Component
    component_table.cell(1, 0).text = "Stress (inverted)"
    component_table.cell(1, 1).text = f"{component_scores['stress_score']:.1f}/100"
    component_table.cell(1, 2).text = "40%"
    
    # Eye Contact Component
    component_table.cell(2, 0).text = "Eye Contact Time"
    component_table.cell(2, 1).text = f"{component_scores['eye_contact_score']:.1f}/100"
    component_table.cell(2, 2).text = "30%"
    
    # Emotion Compliance Component
    component_table.cell(3, 0).text = "Emotion Compliance"
    component_table.cell(3, 1).text = f"{component_scores['emotion_compliance_score']:.1f}/100"
    component_table.cell(3, 2).text = "30%"

    # =====================
    # Emotion Range Compliance ONLY
    # =====================
    doc.add_heading("Emotion Range Compliance", level=1)
    compliance_table = doc.add_table(rows=1, cols=4)
    compliance_table.style = 'Table Grid'
    
    hdr_cells = compliance_table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Observed %'
    hdr_cells[2].text = 'Expected Range'
    hdr_cells[3].text = 'Status'
    
    for emo, (val, ok) in evaluation_report.items():
        row_cells = compliance_table.add_row().cells
        row_cells[0].text = emo
        row_cells[1].text = f"{val*100:.1f}%"
        low, high = st.session_state.EXPECTED_RANGES[emo]
        row_cells[2].text = f"{low*100:.0f}–{high*100:.0f}%"
        status_cell = row_cells[3]
        if ok:
            status_cell.text = "✅ Within range"
            for paragraph in status_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0, 128, 0)
        else:
            status_cell.text = "⚠️ Outside range"
            for paragraph in status_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(220, 20, 60)

    # Recommendations
    doc.add_heading("Key Recommendations", level=1)
    if overall_score >= 80:
        doc.add_paragraph("Overall performance was excellent. The candidate demonstrated strong communication skills with appropriate emotional expression and good eye contact.")
    elif overall_score >= 60:
        doc.add_paragraph("Overall performance was adequate. The candidate showed some good qualities but has room for improvement in certain areas.")
    else:
        doc.add_paragraph("Overall performance needs significant improvement. The candidate demonstrated several areas of concern that should be addressed.")
    
    recommendations = doc.add_paragraph()
    recommendations.add_run("Specific Recommendations:\n").bold = True
    
    if stress_avg > 65:
        recommendations.add_run("• High stress levels detected. Consider providing a more comfortable environment or stress management techniques.\n")
    elif stress_avg > 40:
        recommendations.add_run("• Moderate stress levels observed. Candidate may benefit from relaxation techniques before interviews.\n")
    
    if eye_contact_percentage < 60:
        recommendations.add_run("• Eye contact needs improvement. Practice maintaining good eye contact for longer periods during conversations.\n")
    elif eye_contact_percentage < 70:
        recommendations.add_run("• Eye contact is adequate but could be improved. Aim for maintaining good eye contact 70% or more of the time.\n")
    
    if any(not ok for _, ok in evaluation_report.items()):
        recommendations.add_run("• Emotional expression needs adjustment. Work on maintaining neutral to positive expressions during professional interactions.\n")
    
    if overall_score >= 80:
        recommendations.add_run("• Continue current practices as they are yielding excellent results.\n")

    filepath = "Interview_Report.docx"
    doc.save(filepath)
    return filepath

# =========================
# MULTI-PAGE DASHBOARD
# =========================
st.set_page_config(page_title="Interview Analysis Tool", page_icon="🎥", layout="wide")
page = st.sidebar.radio("Navigate", ["Interview Dashboard", "Admin Dashboard"])

# =========================
# PAGE 1: INTERVIEW DASHBOARD
# =========================
if page == "Interview Dashboard":
    st.title("🎥 Interview Dashboard")
    st.caption("Real-time detection of emotions, eye contact, and stress.")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Start Camera", disabled=st.session_state.camera_active):
            st.session_state.camera_active = True
            st.session_state.history.clear()
            st.session_state.eye_contact_history.clear()
            st.session_state.eye_contact_binary_history.clear()
            st.session_state.stress_history.clear()
            st.rerun()
    with col2:
        if st.button("Stop Camera", disabled=not st.session_state.camera_active):
            st.session_state.camera_active = False
            if len(st.session_state.history) > 0:
                emotion_avg = defaultdict(float)
                for row in st.session_state.history:
                    for k, v in row.items():
                        emotion_avg[k] += v
                n = max(1, len(st.session_state.history))
                for k in emotion_avg:
                    emotion_avg[k] /= n
                eye_contact_avg = sum(st.session_state.eye_contact_history)/max(1,len(st.session_state.eye_contact_history))
                eye_contact_percentage = sum(st.session_state.eye_contact_binary_history)/max(1,len(st.session_state.eye_contact_binary_history))
                stress_avg = sum(st.session_state.stress_history)/max(1,len(st.session_state.stress_history))
                stress_label, _ = get_stress_label(stress_avg)
                evaluation_report = evaluate_emotion_distribution(emotion_avg)

                filepath = generate_report(emotion_avg, eye_contact_avg, eye_contact_percentage, stress_avg, stress_label, evaluation_report)
                with open(filepath, "rb") as f:
                    bytes_data = f.read()
                st.download_button(
                    label="⬇️ Interview Report (Auto Download)",
                    data=bytes_data,
                    file_name=filepath,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="auto_download"
                )

    if st.session_state.camera_active:
        cap = cv2.VideoCapture(0)
        if not cap.isOpened():
            st.error("Cannot access webcam.")
            st.session_state.camera_active = False
        else:
            cap.set(cv2.CAP_PROP_FRAME_WIDTH, 640)
            cap.set(cv2.CAP_PROP_FRAME_HEIGHT, 480)
            cap.set(cv2.CAP_PROP_FPS, FPS)

            left_col, right_col = st.columns([2,2])
            with left_col:
                video_placeholder = st.empty()
            with right_col:
                status_placeholder = st.empty()
                stress_placeholder = st.empty()
                chart_placeholder = st.empty()
                details_placeholder = st.empty()
                evaluation_placeholder = st.empty()

            while st.session_state.camera_active:
                ret, frame = cap.read()
                if not ret:
                    st.error("Failed to capture frame")
                    break

                rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                video_placeholder.image(rgb_frame, channels="RGB", width=500)

                try:
                    result = post_to_backend(frame)
                    (emotion_avg, eye_contact_avg, eye_contact_percentage, stress_avg, stress_label, stress_emoji, eye_contact_perf, dominant_emotion) = update_rollups(result)

                    current_emotion = dominant_emotion  # make dominant emotion prominent
                    confidence = result.get("confidence", 0)
                    eye_contact = result.get("eye_contact", 0)
                    
                    # =======================
                    # 1) INSTANTANEOUS EMOTIONS (Current Frame)
                    # =======================
                    instant_emotions = result.get("all_predictions", {}) or {}  # FIX: ensure dict
                    instant_sorted = sorted(instant_emotions.items(), key=lambda kv: kv[1], reverse=True)

                    if len(instant_sorted) > 0:
                        top_emo, top_val = instant_sorted[0]
                        status_placeholder.markdown(
                            f"<h2 style='text-align: center; color: {'green' if top_emo in ['Neutral','Happy'] else 'red'};'>"
                            f"🔥 Dominant Now: {top_emo} ({top_val*100:.1f}%)</h2>",
                            unsafe_allow_html=True
                        )
                        chart_placeholder.bar_chart({k:[v] for k,v in instant_emotions.items()})
                        # Show ranked list
                        instant_text = "📌 **Instantaneous Emotions (Current Frame):**\n"
                        for emo, val in instant_sorted:
                            instant_text += f"- {emo}: {val*100:.1f}%\n"
                        details_placeholder.markdown(instant_text)
                    else:
                        # FIX: handle no predictions safely (no face / model fallback)
                        status_placeholder.markdown(
                            "<h2 style='text-align: center; color: orange;'>No face detected</h2>",
                            unsafe_allow_html=True
                        )
                        chart_placeholder.bar_chart({e:[0.0] for e in EMOTIONS})
                        details_placeholder.markdown("📌 **Instantaneous Emotions (Current Frame):**\n- No data for this frame.")

                    # Current eye contact & stress (always shown)
                    stress_placeholder.markdown(
                        f"<div style='text-align: center;'>"
                        f"<span style='font-size: 1.2em;'>👀 Eye Contact: {eye_contact}% (Current)</span> | "
                        f"<span style='font-size: 1.2em;'>⏱️ Good Eye Contact Time: {eye_contact_percentage:.1f}%</span> | "
                        f"<span style='font-size: 1.2em;'>{stress_emoji} Stress: {stress_label} ({stress_avg:.1f}/100)</span>"
                        f"</div>",
                        unsafe_allow_html=True
                    )

                    # =======================
                    # 2) SMOOTHED AVERAGES (Rolling Window)
                    # =======================
                    sorted_emotion_avg = sorted(emotion_avg.items(), key=lambda kv: kv[1], reverse=True)
                    avg_text = "📊 **Smoothed Averages (15-frame window):**\n" + \
                            "  |  ".join([f"{e}: {round(v*100,1)}%" for e,v in sorted_emotion_avg])
                    avg_text += f"\n\n⏱️ **Good Eye Contact Time:** {eye_contact_percentage:.1f}% ({eye_contact_perf})"
                    avg_text += f"  |  🔎 **Stress:** {stress_avg:.1f}/100 ({stress_label})"
                    evaluation_placeholder.markdown(avg_text)

                    evaluation_report = evaluate_emotion_distribution(emotion_avg)
                    eval_text = "🎯 **Emotion Ranges:**\n"
                    for k, (val, ok) in evaluation_report.items():
                        percent = round(val*100,1)
                        low, high = st.session_state.EXPECTED_RANGES[k]
                        status = "✅" if ok else "⚠️"
                        eval_text += f"- {k}: {percent}% (exp. {int(low*100)}–{int(high*100)}%) {status}\n"
                    evaluation_placeholder.markdown(eval_text)

                except Exception as e:
                    st.error(f"Backend error: {e}")

    else:
        st.info("Click 'Start Camera' to begin detection")

# =========================
# PAGE 2: ADMIN DASHBOARD
# =========================
if page == "Admin Dashboard":
    st.title("⚙️ Admin Dashboard")
    st.caption("Configure weights and expected ranges for interview evaluation.")

    st.subheader("📌 Negative Emotion Weights (for stress calculation)")
    for emo in ['Angry','Sad','Fear','Disgust']:
        col1, col2 = st.columns([4,1])
        with col1:
            st.session_state.EMOTION_WEIGHTS[emo] = st.slider(
                f"{emo} weight", 0.0, 2.0, st.session_state.EMOTION_WEIGHTS.get(emo, IDEAL_EMOTION_WEIGHTS[emo]), 0.1, key=f"emo_{emo}"
            )
        with col2:
            if st.button("Reset to Default ", key=f"reset_{emo}"):
                st.session_state.EMOTION_WEIGHTS[emo] = IDEAL_EMOTION_WEIGHTS[emo]
                save_settings()
                st.rerun()

    st.subheader("📌 Stress Component Weights")
    for comp in ["emotions","eye_contact"]:
        col1, col2 = st.columns([4,1])
        with col1:
            st.session_state.STRESS_WEIGHTS[comp] = st.slider(
                f"{comp} weight", 0.0, 1.0, st.session_state.STRESS_WEIGHTS[comp], 0.05, key=f"comp_{comp}"
            )
        with col2:
            if st.button("Reset to Default ", key=f"reset_{comp}"):
                st.session_state.STRESS_WEIGHTS[comp] = IDEAL_STRESS_WEIGHTS[comp]
                save_settings()
                st.rerun()
    s = sum(st.session_state.STRESS_WEIGHTS.values())
    for k in st.session_state.STRESS_WEIGHTS:
        st.session_state.STRESS_WEIGHTS[k] /= s

    st.subheader("📌 Expected Emotion Ranges")
    for emo in st.session_state.EXPECTED_RANGES.keys():
        col1, col2 = st.columns([4,1])
        with col1:
            low, high = st.session_state.EXPECTED_RANGES[emo]
            st.session_state.EXPECTED_RANGES[emo] = st.slider(
                f"{emo} % range", 0.0, 1.0, (low, high), 0.05, key=f"range_{emo}"
            )
        with col2:
            if st.button("Reset to Default ", key=f"reset_range_{emo}"):
                st.session_state.EXPECTED_RANGES[emo] = IDEAL_EXPECTED_RANGES[emo]
                save_settings()
                st.rerun()

    save_settings()
    st.success("✅ Admin settings saved. They will apply in Interview Dashboard.")
